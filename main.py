import os
import shutil
import time
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse
from pathlib import Path
from docx import Document
from groq import Groq
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

app = FastAPI()

from fastapi.middleware.cors import CORSMiddleware

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins for now. In production, replace ["*"] with your frontend URL.
    allow_credentials=True,
    allow_methods=["*"],  # Allows all HTTP methods (GET, POST, etc.)
    allow_headers=["*"],  # Allows all headers
)

# Setup Directories
BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "input"
OUTPUT_DIR = BASE_DIR / "output"
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

def process_entire_document(full_text, language):
    """Sends the whole document to Groq and requests code in a specific language."""
    try:
        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system", 
                    "content": (
                        f"You are a technical assignment solver specializing in {language}. "
                        "Identify every individual lab task or question in the text provided. "
                        f"For each one, provide the solution ONLY in {language} code. "
                        "Output exactly this format:\n\n"
                        "TASK [Number]\n"
                        "[Code]\n"
                        "---END_OF_TASK---\n\n"
                        "STRICT RULES:\n"
                        f"1. Provide ONLY the raw {language} code. No explanations, no introductions, and no markdown backticks (```).\n"
                        f"2. Ensure all necessary {language} libraries, headers, or boilerplate (like main functions) are included for each task.\n"
                        "3. Do not include comments or conversational filler.\n"
                        "4. Ignore submission instructions, teacher names, or non-task text."
                    )
                },
                {"role": "user", "content": full_text}
            ],
            temperature=0.1,
        )
        return completion.choices[0].message.content.strip()
    except Exception as e:
        return f"Error from AI: {str(e)}"

@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    name: str = Form(...),
    roll_number: str = Form(...),
    section: str = Form(...),
    subject: str = Form(...),
    language: str = Form("Python")  # Default to Python if not specified
):
    if not file.filename.endswith('.docx'):
        raise HTTPException(400, "Only .docx files are supported.")

    # 1. Save and extract raw text
    input_path = UPLOAD_DIR / file.filename
    with input_path.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    doc = Document(input_path)
    full_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

    # 2. Get the entire solution from AI in one shot with the chosen language
    print(f"Processing entire document for {language} solutions...")
    full_solution = process_entire_document(full_text, language)

    # 3. Build the Word Document
    output_doc = Document()
    
    # Header Section
    output_doc.add_heading(f"Subject: {subject}", level=0)
    output_doc.add_paragraph(f"Name: {name} | Roll: {roll_number} | Section: {section} | Language: {language}")
    output_doc.add_paragraph("-" * 60)

    # Split the AI response using the custom separator
    task_solutions = full_solution.split("---END_OF_TASK---")

    for task in task_solutions:
        clean_task = task.strip()
        if len(clean_task) > 20:  # Avoid empty splits
            # Split the 'TASK X' header from the code
            lines = clean_task.split('\n', 1)
            if len(lines) > 1:
                output_doc.add_heading(lines[0], level=1) # The Task Label
                output_doc.add_paragraph(lines[1])        # The actual Code
            else:
                output_doc.add_paragraph(clean_task)
            
            output_doc.add_page_break()

    # 4. Save the file with a timestamp to avoid permission errors
    timestamp = int(time.time())
    output_filename = f"Solved_{timestamp}_{file.filename}"
    output_path = OUTPUT_DIR / output_filename
    output_doc.save(output_path)
    
    return {
        "status": "Success",
        "language_used": language,
        "download_url": f"/download/{output_filename}"
    }

@app.get("/download/{filename}")
async def download_file(filename: str):
    file_path = OUTPUT_DIR / filename
    if file_path.exists():
        return FileResponse(path=file_path, filename=filename)
    raise HTTPException(404, "File not found.")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)